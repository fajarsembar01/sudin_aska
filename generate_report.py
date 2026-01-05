"""Generate Word document report for Portal PANBERSS."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Create document
doc = Document()

# Title
title = doc.add_heading('Laporan Lengkap Fitur Portal PANBERSS', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('(Penilaian Bersolek Sekolah)')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

intro = doc.add_paragraph()
intro.add_run('Baik Pak, berikut rincian link beserta penjelasan fitur-fiturnya:')

# ===== Section 1 =====
doc.add_heading('1. 📊 Statistik BERSOLEK', level=1)

link1 = doc.add_paragraph()
link1.add_run('Link: ').bold = True
link1.add_run('https://admin.sudindikju2.com/portal/admin/stats')

doc.add_paragraph('Deskripsi Fitur:', style='Intense Quote')
doc.add_paragraph('Halaman ini menyajikan ringkasan data penilaian sekolah secara menyeluruh dengan fitur-fitur berikut:')

# Dashboard Statistik
doc.add_heading('Dashboard Statistik:', level=2)
doc.add_paragraph('Total sekolah terdaftar dan yang sudah aktif', style='List Bullet')
doc.add_paragraph('Jumlah penilaian berstatus Draft (belum disubmit) dan Submitted (sudah disubmit)', style='List Bullet')
doc.add_paragraph('Rata-rata skor dari seluruh penilaian yang telah disubmit', style='List Bullet')

# Peta Interaktif
doc.add_heading('Peta Interaktif (Map):', level=2)
doc.add_paragraph('Menampilkan lokasi sekolah dengan warna berbeda berdasarkan skor:', style='List Bullet')
doc.add_paragraph('🟢 Hijau = Skor ≥85 (Baik)', style='List Bullet 2')
doc.add_paragraph('🟡 Kuning = Skor 70-84 (Cukup)', style='List Bullet 2')
doc.add_paragraph('🔴 Merah = Skor 55-69 (Kurang)', style='List Bullet 2')
doc.add_paragraph('⚫ Hitam = Skor <55 (Sangat Kurang)', style='List Bullet 2')
doc.add_paragraph('Terdapat fitur Heatmap (Radius Rawan) yang dapat ditampilkan/sembunyikan untuk mengidentifikasi area dengan konsentrasi sekolah skor rendah', style='List Bullet')

note1 = doc.add_paragraph()
note1.add_run('⚠️ Catatan: ').bold = True
note1.add_run('Seluruh data saat ini bersifat contoh, jadi lokasi titik-titik sekolah pada map tidak akurat. Ini digunakan untuk menguji fitur-fitur yang telah dibuat')

# Grafik
doc.add_heading('Grafik Sebaran Nilai:', level=2)
doc.add_paragraph('Menampilkan distribusi nilai sekolah dalam 9 kategori (<60 sampai 95-100)', style='List Bullet')
doc.add_paragraph('Pie chart status penilaian (Draft vs Submitted)', style='List Bullet')

# Perankingan
doc.add_heading('Perankingan Sekolah:', level=2)
doc.add_paragraph('Tab Terbaik: 5 sekolah dengan skor tertinggi', style='List Bullet')
doc.add_paragraph('Tab Terendah: 5 sekolah dengan skor terendah - untuk fokus perbaikan prioritas', style='List Bullet')

# Galeri
doc.add_heading('Galeri Foto Ruangan:', level=2)
doc.add_paragraph('Dapat diurutkan berdasarkan: Random, Terbaru, atau Terendah', style='List Bullet')
doc.add_paragraph('💡 Filter ke skor terendah untuk memfokuskan ruangan yang paling penting untuk segera ditangani', style='List Bullet')

# Daftar Penilaian
doc.add_heading('Daftar Penilaian Terbaru:', level=2)
doc.add_paragraph('Tabel yang dapat di-sort berdasarkan nama sekolah, staff, skor rata-rata, dan tanggal', style='List Bullet')
doc.add_paragraph('Filter berdasarkan jenjang (SD, SMP, SMA, dll) dan periode', style='List Bullet')
doc.add_paragraph('Aksi: Lihat detail dan hapus penilaian', style='List Bullet')

# Export
doc.add_heading('Export Excel:', level=2)
doc.add_paragraph('Tombol export untuk mengunduh data penilaian ke format Excel', style='List Bullet')

# ===== Section 2 =====
doc.add_page_break()
doc.add_heading('2. ⚙️ Setup Data Ruangan & Aspek BERSOLEK', level=1)

link2 = doc.add_paragraph()
link2.add_run('Link: ').bold = True
link2.add_run('https://admin.sudindikju2.com/portal/admin/setup')

doc.add_paragraph('Deskripsi Fitur:', style='Intense Quote')
doc.add_paragraph('Halaman ini digunakan untuk konfigurasi master data dengan fitur-fitur:')

doc.add_heading('Tambah Sekolah:', level=2)
doc.add_paragraph('Input NPSN, nama, jenjang, alamat, kelurahan, dan status (Negeri/Swasta)', style='List Bullet')

doc.add_heading('Tambah Ruangan:', level=2)
doc.add_paragraph('Input nama ruangan, urutan, deskripsi, dan kategori (Umum/Akademik/Fasilitas/Sanitasi)', style='List Bullet')

doc.add_heading('Tambah Aspek Penilaian:', level=2)
doc.add_paragraph('Pilih ruangan, masukkan nama aspek, urutan, dan deskripsi', style='List Bullet')

doc.add_heading('Preview Konfigurasi per Jenjang:', level=2)
doc.add_paragraph('Tab SD: Kelas 1-6 + ruang lainnya', style='List Bullet')
doc.add_paragraph('Tab SMP: Kelas 7-9 + ruang lainnya', style='List Bullet')
doc.add_paragraph('Tab SMA/SMK: Kelas 10-12 + ruang khusus (Lab Komputer, Lab IPA/Kimia)', style='List Bullet')

doc.add_heading('Daftar Ruangan & Aspek:', level=2)
doc.add_paragraph('Tabel dengan kolom: Nama Ruangan, Kategori, Jumlah Aspek, Status (Aktif/Nonaktif)', style='List Bullet')
doc.add_paragraph('Expandable detail untuk edit ruangan dan menambah/edit aspek langsung (inline)', style='List Bullet')
doc.add_paragraph('Edit langsung nama, kategori, urutan, status, dan deskripsi ruangan', style='List Bullet')
doc.add_paragraph('Hapus ruangan (jika ada penilaian menggunakan ruangan ini, ruangan akan dinonaktifkan saja)', style='List Bullet')

doc.add_heading('Daftar Sekolah:', level=2)
doc.add_paragraph('Menampilkan 20 sekolah pertama dengan info NPSN, Nama, Jenjang, Status, dan Kecamatan', style='List Bullet')

# ===== Section 3 =====
doc.add_page_break()
doc.add_heading('3. 📱 Akses Tim Monev BERSOLEK', level=1)

link3 = doc.add_paragraph()
link3.add_run('Link: ').bold = True
link3.add_run('https://admin.sudindikju2.com/portal/schools')

doc.add_paragraph('Deskripsi Fitur:', style='Intense Quote')
doc.add_paragraph('Halaman ini digunakan oleh tim monitoring dan evaluasi yang ditugaskan ke sekolah:')

doc.add_heading('Pilih Sekolah:', level=2)
doc.add_paragraph('Daftar sekolah yang tersedia untuk penilaian', style='List Bullet')

note2 = doc.add_paragraph()
note2.add_run('⚠️ Catatan: ').bold = True
note2.add_run('Saat ini data yang tersedia hanya SDN Semper Barat 01 atau NPSN: 20100682')

doc.add_heading('Fitur Penilaian Langsung:', level=2)
doc.add_paragraph('Tim monev bisa langsung ambil foto ruangan', style='List Bullet')
doc.add_paragraph('Foto sudah dilengkapi dengan timestamp (waktu pengambilan)', style='List Bullet')
doc.add_paragraph('Foto sudah dilengkapi dengan lokasi GPS (latitude, longitude)', style='List Bullet')

doc.add_heading('Konfigurasi Ruangan Sekolah:', level=2)
doc.add_paragraph('Sekolah dapat memilih ruangan mana saja yang dimiliki', style='List Bullet')
doc.add_paragraph('Dikelompokkan per jenjang: Umum, SD, SMP, SMA/SMK', style='List Bullet')
doc.add_paragraph('Setiap ruangan memiliki aspek penilaian yang dapat dilihat', style='List Bullet')

doc.add_heading('Kelengkapan Profil Sekolah:', level=2)
doc.add_paragraph('Modal wizard 4 langkah untuk melengkapi data sekolah:', style='List Bullet')
doc.add_paragraph('1. Lokasi & GMaps (Alamat, Kelurahan, Link GMaps)', style='List Bullet 2')
doc.add_paragraph('2. Kapasitas (Jumlah Siswa, Bangku Kosong, Guru, Tendik, Rombel)', style='List Bullet 2')
doc.add_paragraph('3. Kontak Sekolah (Telepon, Operator, Email CS)', style='List Bullet 2')
doc.add_paragraph('4. Channel Opsional (Instagram, TikTok, YouTube, WA Channel)', style='List Bullet 2')

# ===== Credentials Table =====
doc.add_page_break()
doc.add_heading('🔐 Kredensial Akses:', level=1)

table = doc.add_table(rows=2, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Email'
hdr_cells[1].text = 'Password'
hdr_cells[2].text = 'Keterangan'
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True

# Data
row_cells = table.rows[1].cells
row_cells[0].text = 'admin@sekolah.sch.id'
row_cells[1].text = '12345678'
row_cells[2].text = 'Akses untuk semua halaman'

# ===== Notes =====
doc.add_heading('📌 Catatan Penting:', level=1)

doc.add_paragraph('Data Contoh: Seluruh data penilaian saat ini bersifat contoh/dummy untuk testing fitur', style='List Number')
doc.add_paragraph('Lokasi Map: Koordinat sekolah pada peta tidak akurat karena data masih sampel', style='List Number')
doc.add_paragraph('Periode Penilaian: Sistem mendukung multiple periode penilaian yang dapat difilter', style='List Number')

# Score System
score_para = doc.add_paragraph(style='List Number')
score_para.add_run('Sistem Skor: ').bold = True
score_para.add_run('Skor dinilai dalam skala 0-3 per aspek:')

# Score table
score_table = doc.add_table(rows=5, cols=2)
score_table.style = 'Table Grid'
score_table.alignment = WD_TABLE_ALIGNMENT.CENTER

score_hdr = score_table.rows[0].cells
score_hdr[0].text = 'Skor'
score_hdr[1].text = 'Keterangan'
for cell in score_hdr:
    cell.paragraphs[0].runs[0].bold = True

scores = [('0', 'Buruk'), ('1', 'Kurang'), ('2', 'Cukup'), ('3', 'Baik (nilai default jika belum dinilai)')]
for i, (skor, ket) in enumerate(scores, 1):
    score_table.rows[i].cells[0].text = skor
    score_table.rows[i].cells[1].text = ket

doc.add_paragraph()
formula = doc.add_paragraph()
formula.add_run('Skor kemudian dikonversi ke persentase 0-100% dengan rumus: ').italic = True
formula.add_run('(skor / 3) × 100').bold = True

# Closing
doc.add_paragraph()
closing = doc.add_paragraph()
closing.add_run('Demikian laporan fitur Portal PANBERSS yang sudah dikembangkan. Jika ada pertanyaan atau membutuhkan penjelasan lebih lanjut terkait fitur-fitur tersebut, silakan hubungi kami. 🙏')

# Save
output_path = '/Users/ainunfajar/SUDIN_ASKA/ai-agent-sekolah/Laporan_Portal_PANBERSS.docx'
doc.save(output_path)
print(f"✅ Dokumen berhasil disimpan: {output_path}")
